"""
Fix the v2 manuscript:
  1. Remove the stale duplicate deployments table (second w:tbl in the document).
  2. Replace the Section 10 Configuration text paragraphs with a real Word table.
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = (r"C:\Users\busta\Code\GaInSn-CALPHAD\cooling\Energy-submission"
        r"\Cooling_Energy_manuscript_v2.docx")

doc = Document(PATH)

# ---------------------------------------------------------------------------
# 1. Remove the duplicate deployments table (the one with "Biforest" / typos).
#    It is the docx TABLE 2 — identified by the cell text "Biforest".
# ---------------------------------------------------------------------------
removed = 0
for tbl in doc.tables:
    cell_texts = [c.text for row in tbl.rows for c in row.cells]
    if any('Biforest' in t or 'Project Mining' in t for t in cell_texts):
        tbl._element.getparent().remove(tbl._element)
        removed += 1
        print(f"Removed duplicate deployments table ({removed}).")
        break

if not removed:
    print("WARNING: duplicate table not found — check identifiers.")

# ---------------------------------------------------------------------------
# 2. Build the real Table 4 and replace the Configuration text paragraphs.
#
#    Row data: (config label, k, q_max, T_out, Ex, refs)
# ---------------------------------------------------------------------------
ROWS = [
    ("Forced-air convection",
     "0.026", "~1", "25-35", "~8%", "[5]"),
    ("Liquid cold plate\n(single-phase water)",
     "0.6", "~30", "50-65", "12-15%", "[5]"),
    ("Dielectric single-phase immersion",
     "0.06-0.14", "~65", "65-75", "16-19%", "[10, 11, 12]"),
    ("Dielectric two-phase immersion",
     "0.06-0.14", "~200", "50-65", "12-16%", "[15, 38]"),
    ("Gallium-based liquid metal\n(single-phase)",
     "16-86", "~500", ">80", "20-23%", "[22, 27, 31, 33]"),
    ("Embedded single-phase microfluidics\n(monolithic Cu)",
     "~400", "790-900", "50-70", "12-16%", "[32]"),
    ("Embedded two-phase\n(jet / microchannels)",
     "--", "~3,000", "50-60", "12-14%", "[38]"),
]

HEADERS = [
    "Configuration",
    "k\n(W m⁻¹ K⁻¹)",
    "qₘₐˣ\n(W cm⁻²)",
    "T₀ᵁᵗ\n(°C)",
    "Ex",
    "Refs",
]

# Column widths in twentieths of a point (1440 = 1 inch)
# Total content width ~6.0 inches = 8640 twips
COL_WIDTHS = [2160, 1080, 1080, 1000, 800, 520]   # sum = 6640


def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')


def cell_shade(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def set_thin_borders(tbl):
    """Apply thin single borders to every cell."""
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'AAAAAA')
                tcBorders.append(b)
            tcPr.append(tcBorders)


def write_cell(cell, text, bold=False, center=False, size=9):
    for para in cell.paragraphs:
        para._element.getparent().remove(para._element)
    for line in text.split('\n'):
        p = cell.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)


# Build the table
tbl = doc.add_table(rows=1 + len(ROWS), cols=len(HEADERS))
tbl.style = 'Table Grid'

# Header row
hrow = tbl.rows[0]
for j, (hdr, w) in enumerate(zip(HEADERS, COL_WIDTHS)):
    cell = hrow.cells[j]
    set_cell_width(cell, w)
    set_cell_margins(cell)
    cell_shade(cell, '2E4057')       # dark slate header
    write_cell(cell, hdr, bold=True, center=(j > 0), size=8)
    # White text
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Data rows
for i, (config, k, qmax, tout, ex, refs) in enumerate(ROWS):
    row = tbl.rows[i + 1]
    fill = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
    data = [config, k, qmax, tout, ex, refs]
    for j, (val, w) in enumerate(zip(data, COL_WIDTHS)):
        cell = row.cells[j]
        set_cell_width(cell, w)
        set_cell_margins(cell)
        cell_shade(cell, fill)
        write_cell(cell, val, bold=False, center=(j > 0), size=9)

set_thin_borders(tbl)

# The table was appended at the end of the document body — grab its element.
tbl_elem = tbl._element
tbl_elem.getparent().remove(tbl_elem)   # detach; we will re-insert in place

# ---------------------------------------------------------------------------
# 3. Find the Configuration paragraphs in Section 10 and replace them with
#    the real table element, inserted after the caption paragraph.
# ---------------------------------------------------------------------------
paras = doc.paragraphs
cfg_indices = []
caption_idx = None

for idx, para in enumerate(paras):
    t = para.text.strip()
    if t.startswith('Table 4.'):
        caption_idx = idx
    if t.startswith('Configuration ') and ' - ' in t:
        cfg_indices.append(idx)

print(f"Caption paragraph index: {caption_idx}")
print(f"Configuration paragraphs: indices {cfg_indices}")

if not cfg_indices:
    raise RuntimeError("Could not find Configuration paragraphs in Section 10.")

# Remove Configuration paragraphs (highest index first to avoid index shift)
body = doc.element.body
cfg_elems = [paras[i]._element for i in cfg_indices]
for elem in cfg_elems:
    elem.getparent().remove(elem)
    print("  Removed a Configuration paragraph.")

# Insert the table after the caption paragraph
caption_elem = paras[caption_idx]._element
caption_parent = caption_elem.getparent()
cap_pos = list(caption_parent).index(caption_elem)
caption_parent.insert(cap_pos + 1, tbl_elem)
print("Table 4 inserted after caption.")

# ---------------------------------------------------------------------------
# 4. Save
# ---------------------------------------------------------------------------
import os
doc.save(PATH)
print(f"\nSaved: {PATH}  ({os.path.getsize(PATH):,} bytes)")
