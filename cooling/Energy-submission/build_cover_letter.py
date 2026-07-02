"""Build Cover_Letter_Energy.docx for the Energy submission."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = (r"C:\Users\busta\Code\GaInSn-CALPHAD\cooling\Energy-submission"
       r"\Cover_Letter_Energy.docx")

doc = Document()

# --- Page margins (1 inch all around) ---
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def normal(text, bold=False, italic=False, size=11, space_after=0, space_before=0,
           align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def blank(space_after=4):
    normal("", size=11, space_after=space_after)

# --- Sender block ---
normal("Michael E. Bustamante", bold=True)
normal("Odinzen LLC, Houston, TX, United States")
normal("michaelbusta@odinzen.io")
normal("ORCID: 0009-0009-9001-8151")

blank(space_after=8)

normal("[DATE]", italic=True, color=(120, 120, 120))

blank(space_after=8)

normal("The Editors")
normal("Energy")
normal("Elsevier")

blank(space_after=8)

# --- Re line ---
p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(0)
p.paragraph_format.space_before = Pt(0)
p.add_run("Re: ").bold = True
run_re = p.runs[0]
run_re.font.size = Pt(11)
p.add_run(
    "Submission of original manuscript — “Materials Design for Data Center Thermal "
    "Management and Waste Heat Recovery: Dielectric Fluids, Gallium Liquid Metals, and "
    "Embedded Two-Phase Cooling”"
).font.size = Pt(11)

blank(space_after=8)

# --- Salutation ---
normal("Dear Editors,")

blank(space_after=4)

# --- Opening paragraph ---
normal(
    "We are pleased to submit the above manuscript for consideration as a review article "
    "in Energy. The work treats the management of data center waste heat as a materials "
    "design problem and shows that the choice of coolant is an energy systems decision: "
    "the coolant fixes the thermal pathway, which fixes the output temperature and "
    "resistance, which fixes the grade of recoverable heat. As AI and large language "
    "model data centers push global compute electricity demand past 1,000 TWh per year, "
    "that choice determines whether the thermal output is an environmental cost or an "
    "energy asset.",
    space_after=8
)

# --- Contributions header ---
normal("The manuscript makes four contributions that fit Energy’s scope:", space_after=4)

# --- Contributions as indented paragraphs ---
contributions = [
    (
        "A process–structure–property–performance treatment of three coolant "
        "generations (dielectric immersion fluids, gallium-based liquid metals, and embedded "
        "two-phase microfluidics), tied throughout to the grade and recoverable exergy of the "
        "resulting waste heat, with six documented operational heat-reuse deployments and a "
        "lifecycle CO₂ displacement analysis."
    ),
    (
        "A cross-generation quantitative comparison (Section 10, Table 4) spanning seven "
        "representative configurations from forced-air convection (~1 W/cm²) through "
        "embedded two-phase jet microchannels (~3,000 W/cm²), with exergy fractions "
        "computed at a common reference temperature so the grade-versus-flux trade-off is "
        "directly legible."
    ),
    (
        "An original computational result: using an open CALPHAD workflow, the Ga–In–Sn "
        "liquidus is computed from assessed binary descriptions plus a single ternary term "
        "fitted to the measured eutectic, reproducing the Galinstan melting point of "
        "10.7 ± 0.3 °C to within experimental error. This establishes the coolant "
        "melting range as a designable property rather than a catalogue choice, directly "
        "relevant to matching coolant freezing behaviour to district-heating supply windows."
    ),
    (
        "A quantitative account of limits: the exergy already destroyed at the rack, the "
        "grid-intensity breakeven below which heat reuse is net-positive on CO₂, the "
        "corrosion and supply-chain constraints on gallium, and the gap between laboratory "
        "demonstrations and facility-scale deployment."
    ),
]

for i, text in enumerate(contributions, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)

blank(space_after=4)

# --- Prior work anchor ---
normal(
    "The corresponding author’s prior peer-reviewed work provides the experimental "
    "thermodynamic anchor for the gallium analysis: Bustamante, Lilova, Navrotsky, Harvey, "
    "Oishi, “Enthalpies of mixing for alloys liquid below room temperature determined "
    "by oxidative solution calorimetry,” J. Therm. Anal. Calorim. 149 (2024) 4817–4826 "
    "(doi:10.1007/s10973-024-13035-5).",
    space_after=8
)

# --- Declarations ---
normal(
    "The manuscript is original, has not been published in whole or in substantive part "
    "elsewhere, and is not under consideration by any other journal. There are no competing "
    "interests and no external funding. Use of AI-assisted tools (editorial revision, "
    "structural reorganisation, figure-production scripting, formatting) is disclosed in the "
    "manuscript; all factual claims, numerical values, and citations were verified by the "
    "authors against primary sources, and the reference list was checked field-by-field "
    "against resolved Crossref records.",
    space_after=8
)

# --- Closing ---
normal(
    "Thank you for considering this work. We would welcome review and stand ready to respond "
    "promptly to any comments.",
    space_after=12
)

normal("Yours sincerely,")
blank(space_after=12)

normal("Michael E. Bustamante", bold=True)
normal("Odinzen LLC")
normal("Corresponding author")

blank(space_after=6)

normal("Kristina Lilova", bold=True)
normal("Center for Materials of the Universe")
normal("School of Molecular Sciences, Arizona State University")

import os
doc.save(OUT)
print(f"Saved: {OUT}  ({os.path.getsize(OUT):,} bytes)")
