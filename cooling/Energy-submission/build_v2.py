"""Build Cooling_Energy_manuscript_v2.docx.

Steps:
1. Renumber all in-text citations to strict first-appearance (citation) order.
2. Reorder the reference list to match.
3. Insert Section 10 (cross-generation comparison) before the reference list.
"""
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"C:\Users\busta\Downloads\Cooling_Energy_paper_2026-06-28.docx"
OUTPUT = (r"C:\Users\busta\Code\GaInSn-CALPHAD\cooling\Energy-submission"
          r"\Cooling_Energy_manuscript_v2.docx")

doc = Document(INPUT)

# ---------------------------------------------------------------------------
# Citation patterns  - single [9] and multi-citation [9, 10, 11] groups
# ---------------------------------------------------------------------------
CITE_GROUP = re.compile(r'\[(\d+(?:,\s*\d+)*)\]')


def nums_in(text: str) -> list[int]:
    result = []
    for m in CITE_GROUP.finditer(text):
        for part in m.group(1).split(','):
            result.append(int(part.strip()))
    return result


# ---------------------------------------------------------------------------
# 1. Build citation-order map
#    Use para.text (concatenated across all runs) for reliable detection.
# ---------------------------------------------------------------------------
first_app: dict[int, int] = {}
for i, para in enumerate(doc.paragraphs):
    for n in nums_in(para.text):
        if n not in first_app:
            first_app[n] = i

citation_seq = [n for n, _ in sorted(first_app.items(), key=lambda kv: kv[1])]
old_to_new: dict[int, int] = {old: idx + 1 for idx, old in enumerate(citation_seq)}

print(f"Distinct references: {len(old_to_new)}")
for old, new in sorted(old_to_new.items()):
    print(f"  old[{old:>2}] -> new[{new:>2}]")


# ---------------------------------------------------------------------------
# 2. Replace citations in every paragraph.
#
#    Per-run replacement is unreliable when "[n]" spans run boundaries.
#    Strategy: try per-run first; if the full paragraph text still contains
#    unprocessed citations (because some were split), fall back to a
#    full-paragraph text replacement (first run gets all text, rest cleared).
# ---------------------------------------------------------------------------

def encode(m: re.Match) -> str:
    key = m.group(1).replace(' ', '').replace(',', '_')
    return f'[XXOLD_{key}_XX]'


TEMP = re.compile(r'\[XXOLD_([\d_]+)_XX\]')


def decode(m: re.Match) -> str:
    parts = [int(p) for p in m.group(1).split('_')]
    return '[' + ', '.join(str(old_to_new.get(p, p)) for p in parts) + ']'


def replace_para(para) -> None:
    """Renumber citations in a paragraph, handling split-run cases."""
    if not CITE_GROUP.search(para.text):
        return

    # Pass A: try per-run encoding
    for run in para.runs:
        run.text = CITE_GROUP.sub(encode, run.text)

    # Check if any citations remain unencoded (split-run case)
    if CITE_GROUP.search(para.text):
        # Fall back: full paragraph text replacement, clears per-run formatting
        full = CITE_GROUP.sub(encode, para.text)
        full = TEMP.sub(decode, full)
        if para.runs:
            para.runs[0].text = full
            for run in para.runs[1:]:
                run.text = ""
        return

    # Pass B: decode encoded markers
    for run in para.runs:
        run.text = TEMP.sub(decode, run.text)


for para in doc.paragraphs:
    replace_para(para)

print("Citation renumbering applied.")

# ---------------------------------------------------------------------------
# 3. Identify reference list entries and reorder by new number
# ---------------------------------------------------------------------------
REF_ENTRY = re.compile(r'^\[(\d+)\]')

ref_paras: list[tuple[int, object]] = []
for para in doc.paragraphs:
    m = REF_ENTRY.match(para.text.strip())
    if m:
        ref_paras.append((int(m.group(1)), para._element))

print(f"\nReference list entries ({len(ref_paras)} found):")
for num, _ in sorted(ref_paras, key=lambda x: x[0]):
    print(f"  [{num}]")

ref_paras.sort(key=lambda x: x[0])

if ref_paras:
    first_ref_elem = ref_paras[0][1]
    parent = first_ref_elem.getparent()
    insert_idx = list(parent).index(first_ref_elem)
    for _, elem in ref_paras:
        parent.remove(elem)
    for offset, (_, elem) in enumerate(ref_paras):
        parent.insert(insert_idx + offset, elem)
    print("Reference list reordered.")

# ---------------------------------------------------------------------------
# 4. Insert Section 10 before the reference list
#
# All citation numbers here use the NEW numbering (from the map above):
#   new[5]  = Euroheat & Power (district heating 60-80 C)
#   new[10, 11, 12] = 3M / Shell / Engineered Fluids TDS
#   new[13] = ECHA PFAS restriction
#   new[14] = EU Energy Efficiency Directive 2023/1791
#   new[15] = Birbarah et al. 2020 IJHMT (water immersion)
#   new[22] = Yan et al. 2024 IJHMT (Ga-In-Sn TIM performance)
#   new[24] = Bustamante & Lilova (companion CALPHAD assessment)
#   new[27] = Ma & Liu 2007 Phys Lett A (nano liquid-metal coolant)
#   new[29] = Yang et al. 2018 Appl Phys Lett (liquid-metal thermal switch)
#   new[30] = Yang et al. 2019 IEEE T-CPMT
#   new[31] = Deng & Liu 2009 Appl Phys A (gallium corrosion)
#   new[32] = Gebrael et al. 2022 Nature Electronics (monolithic Cu)
#   new[33] = Zhang et al. 2020 IJHMT (supercooling nucleation agents)
#   new[38] = Bar-Cohen et al. 2021 IEEE T-CPMT (ICECool evaporative)
# ---------------------------------------------------------------------------
SEC10: list[tuple[str, str]] = [
    ("Heading2",
     "10. Cross-Generation Quantitative Comparison"),
    ("Normal",
     "The three cooling generations span four orders of magnitude in peak areal "
     "heat flux and roughly double the fraction of input electricity recoverable "
     "as useful exergy. Table 4 summarises seven representative configurations "
     "from forced-air rejection through embedded two-phase microfluidics, drawing "
     "on data from Sections 2-9 and the primary sources cited there. Exergy "
     "fractions are computed at T0 = 283 K (10 degrees C) as "
     "Ex/W = 1 - T0/Tout, where Tout is the absolute outlet temperature."),
    ("Normal",
     "Table 4. Cross-generation thermal performance. "
     "k: coolant thermal conductivity (W m-1 K-1). "
     "q_max: peak demonstrated areal heat flux (W cm-2). "
     "Tout: coolant outlet temperature (degrees C). "
     "Ex: recoverable exergy fraction of facility input at Tout "
     "(T0 = 283 K, 10 degrees C)."),
    ("Normal",
     "Configuration 1 - Forced-air convection: "
     "k = 0.026 W m-1 K-1, q_max ~1 W cm-2, "
     "Tout = 25-35 degrees C, Ex ~8%. "
     "Below the 60-80 degrees C district heating supply temperature [5]; "
     "no heat-pump-free waste heat reuse path."),
    ("Normal",
     "Configuration 2 - Liquid cold plate (single-phase water): "
     "k = 0.6 W m-1 K-1, q_max ~30 W cm-2, "
     "Tout = 50-65 degrees C, Ex ~12-15%. "
     "Dominant architecture in 2025 hyperscale AI deployments; "
     "typically requires a small heat-pump lift to reach district heating supply "
     "temperatures [5]."),
    ("Normal",
     "Configuration 3 - Dielectric single-phase immersion: "
     "k = 0.06-0.14 W m-1 K-1 [10, 11, 12], "
     "q_max ~65 W cm-2, "
     "Tout = 65-75 degrees C, Ex ~16-19%. "
     "First configuration to cross the district heating threshold without a heat "
     "pump; six operational deployments documented in Section 4."),
    ("Normal",
     "Configuration 4 - Dielectric two-phase immersion: "
     "k = 0.06-0.14 W m-1 K-1, "
     "q_max up to ~200 W cm-2, "
     "Tout = 50-65 degrees C, Ex ~12-16% [38, 15]. "
     "Higher flux than single-phase but exit temperature anchored near the "
     "refrigerant saturation point; PFAS restrictions limit fluid selection [13, 14]."),
    ("Normal",
     "Configuration 5 - Gallium-based liquid metal (single-phase): "
     "k = 16-86 W m-1 K-1 [22, 27], "
     "q_max ~500 W cm-2, "
     "Tout > 80 degrees C, Ex ~20-23%. "
     "The only generation that simultaneously raises both flux capacity and waste "
     "heat grade; corrosion [31] and supercooling [33] require active management."),
    ("Normal",
     "Configuration 6 - Embedded single-phase microfluidics (monolithic Cu): "
     "k ~400 W m-1 K-1 (copper), "
     "q_max ~790-900 W cm-2, "
     "Tout = 50-70 degrees C, Ex ~12-16% [32]. "
     "Thermal removal directly from the die surface demonstrated in "
     "production-grade silicon packages [32]."),
    ("Normal",
     "Configuration 7 - Embedded two-phase (jet impingement / microchannels): "
     "q_max up to 3000 W cm-2, "
     "Tout = 50-60 degrees C, Ex ~12-14% [38]. "
     "Highest demonstrated heat flux; exit temperature anchored near the "
     "refrigerant saturation point, limiting waste heat grade."),
    ("Normal",
     "The comparison in Table 4 quantifies the core trade-off: configurations "
     "optimised for maximum heat flux deliver the lowest recoverable exergy "
     "fraction, while gallium-based liquid metals uniquely advance both metrics. "
     "The companion thermodynamic assessment [24] maps the composition space "
     "that controls the operational window of Configuration 5, supporting the "
     "transition from a fixed-alloy product to a designable coolant [22, 29, 30]."),
]


def make_para(text: str, style: str) -> object:
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


# Find the "References" heading paragraph - insert Section 10 before it,
# so Section 10 stays in the body (not inside the reference section).
# Fall back to the first [n] ref list paragraph if no heading found.
first_ref_anchor = None
for para in doc.paragraphs:
    t = para.text.strip()
    if t.lower() in ("references", "bibliography", "works cited", "literature cited"):
        first_ref_anchor = para._element
        break

if first_ref_anchor is None:
    for para in doc.paragraphs:
        if REF_ENTRY.match(para.text.strip()):
            first_ref_anchor = para._element
            break

if first_ref_anchor is None:
    raise RuntimeError("Could not locate References heading or list start")

# Insert Section 10 before the References heading using direct index insertion.
# addprevious() in a loop inserts each new element just before the anchor,
# reversing the order. Using parent.insert() at a fixed index keeps order correct.
ref_parent = first_ref_anchor.getparent()
ref_idx = list(ref_parent).index(first_ref_anchor)

# Build the full insertion list: [blank, heading, para1, ..., para10, blank]
inserts = [make_para("", "Normal")]          # blank before heading
for style, text in SEC10:
    inserts.append(make_para(text, style))
inserts.append(make_para("", "Normal"))      # blank after last para

for offset, elem in enumerate(inserts):
    ref_parent.insert(ref_idx + offset, elem)

print("Section 10 inserted.")

# ---------------------------------------------------------------------------
# 5. Save
# ---------------------------------------------------------------------------
doc.save(OUTPUT)
import os
print(f"\nSaved: {OUTPUT}  ({os.path.getsize(OUTPUT):,} bytes)")
