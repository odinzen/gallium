"""Swap citations [8] and [9] in manuscript v2 (text + reference list)."""
import re
from docx import Document

INPUT  = OUTPUT = (r"C:\Users\busta\Code\GaInSn-CALPHAD\cooling\Energy-submission"
                   r"\Cooling_Energy_manuscript_v2.docx")

doc = Document(INPUT)

# Three-step swap via a sentinel to avoid double-replacement.
# Matches [8] and [9] as whole citation brackets, not inside larger numbers.
CITE = re.compile(r'\[(\d+(?:,\s*\d+)*)\]')

def swap_nums(text):
    def replace(m):
        parts = [p.strip() for p in m.group(1).split(',')]
        out = []
        for p in parts:
            if p == '8':
                out.append('__NINE__')
            elif p == '9':
                out.append('__EIGHT__')
            else:
                out.append(p)
        return '[' + ', '.join(out) + ']'
    text = CITE.sub(replace, text)
    text = text.replace('__NINE__', '9').replace('__EIGHT__', '8')
    return text

def swap_para(para):
    if '[8]' not in para.text and '[9]' not in para.text:
        return
    # Per-run pass
    for run in para.runs:
        if '[8]' in run.text or '[9]' in run.text:
            run.text = swap_nums(run.text)
    # Fall back if split-run left markers in para.text
    if '[8]' in para.text or '[9]' in para.text:
        full = swap_nums(para.text)
        if para.runs:
            para.runs[0].text = full
            for run in para.runs[1:]:
                run.text = ''

for para in doc.paragraphs:
    swap_para(para)

# Swap reference list entries [8] and [9] in place
REF_ENTRY = re.compile(r'^\[(\d+)\]')
ref8 = ref9 = None
for para in doc.paragraphs:
    m = REF_ENTRY.match(para.text.strip())
    if m:
        n = int(m.group(1))
        if n == 8:
            ref8 = para
        elif n == 9:
            ref9 = para

if ref8 and ref9:
    # Keep labels [8] and [9] in their positions; swap only the body text.
    body8 = ref8.text[ref8.text.index(']') + 1:]  # text after "[8]"
    body9 = ref9.text[ref9.text.index(']') + 1:]  # text after "[9]"
    new8 = '[8]' + body9
    new9 = '[9]' + body8
    if ref8.runs:
        ref8.runs[0].text = new8
        for r in ref8.runs[1:]: r.text = ''
    if ref9.runs:
        ref9.runs[0].text = new9
        for r in ref9.runs[1:]: r.text = ''
    print(f"Swapped ref list body text (labels stay in place).")
    print(f"  [8] now: {new8[:80]}")
    print(f"  [9] now: {new9[:80]}")
else:
    print(f"WARNING: could not find both ref list entries (ref8={ref8}, ref9={ref9})")

import os
doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}  ({os.path.getsize(OUTPUT):,} bytes)")
